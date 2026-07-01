#!/usr/bin/env python3
"""Generate bilingual 参考消息 Word doc — v2 redesigned"""
import json, os, re, sys, time, urllib.request, ssl
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

ctx = ssl.create_default_context()
ctx.check_hostname = False; ctx.verify_mode = ssl.CERT_NONE

SRC_META = {
    "BBC":("C41E3A","GBR BBC News"),"Reuters":("FF8000","GBR Reuters"),
    "Bloomberg":("0053A0","USA Bloomberg"),"AP":("1A1A1A","USA Associated Press"),
    "AFP":("003D7A","FRA AFP"),"Nikkei Asia":("8B0000","JPN Nikkei Asia"),
    "APP (Pakistan)":("014B1A","PAK APP Pakistan"),"SAnews (S.Africa)":("006633","ZAF SAnews"),
}
SRC_FLAGS = {"BBC":"GBR","Reuters":"GBR","Bloomberg":"USA","AP":"USA","AFP":"FRA",
             "Nikkei Asia":"JPN","APP (Pakistan)":"PAK","SAnews (S.Africa)":"ZAF"}

def add_divider(doc, color="C41E3A", width_pt=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(int(width_pt * 8)))
    bottom.set(qn('w:color'), color)
    bottom.set(qn('w:space'), '4')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_section_header(doc, text, color):
    add_divider(doc, color, 10)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(14); r.bold = True
    r.font.color.rgb = RGBColor.from_string(color)

def translate(texts, batch_size=20):
    """Translate titles. Uses DeepSeek API if available, else GLM."""
    api_key = os.environ.get("DEEPSEEK_API_KEY","") or os.environ.get("GLM_API_KEY","")
    if not api_key or api_key == "***":
        return [""]*len(texts)
    
    is_deepseek = bool(os.environ.get("DEEPSEEK_API_KEY",""))
    if is_deepseek:
        api_url = "https://api.deepseek.com/v1/chat/completions"
        model = "deepseek-chat"
    else:
        api_url = "https://open.bigmodel.cn/api/paas/v4/chat/completions"
        model = "glm-4.7-flash"
    
    results = []
    for i in range(0, len(texts), batch_size):
        chunk = texts[i:i+batch_size]
        prompt = "Translate these English news headlines to Chinese. One per line, no numbering:\n" + "\n".join(f"- {t}" for t in chunk)
        try:
            payload = json.dumps({"model":model,"messages":[{"role":"user","content":prompt}],"max_tokens":2048}).encode()
            req = urllib.request.Request(api_url, data=payload,
                headers={"Authorization":f"Bearer {api_key}","Content-Type":"application/json"})
            resp = json.loads(urllib.request.urlopen(req, timeout=60, context=ctx).read())
            text = resp["choices"][0]["message"]["content"].strip()
            lines = [l.strip().lstrip("- ").strip('"').strip("'") for l in text.split("\n") if l.strip()]
            results.extend(lines[:len(chunk)])
        except:
            results.extend([""]*len(chunk))
        time.sleep(0.5)
    return results

def translate_full(articles):
    """Translate full_text for articles that have it."""
    api_key = os.environ.get("DEEPSEEK_API_KEY","") or os.environ.get("GLM_API_KEY","")
    if not api_key or api_key == "***":
        return {}
    
    is_deepseek = bool(os.environ.get("DEEPSEEK_API_KEY",""))
    api_url = "https://api.deepseek.com/v1/chat/completions" if is_deepseek else "https://open.bigmodel.cn/api/paas/v4/chat/completions"
    model = "deepseek-chat" if is_deepseek else "glm-4.7-flash"
    
    to_translate = [(id(a), a.get("full_text","")[:800]) for a in articles if a.get("full_text") and len(a["full_text"])>80]
    if not to_translate: return {}
    
    results = {}
    sep = "\n===NEXT===\n"
    for i in range(0, len(to_translate), 4):
        chunk = to_translate[i:i+4]
        combined = sep.join([t[1] for t in chunk])
        prompt = "Translate these English news excerpts to Chinese. Keep '===NEXT==='. Return ONLY Chinese:\n\n" + combined
        try:
            payload = json.dumps({"model":model,"messages":[{"role":"user","content":prompt}],"max_tokens":4096}).encode()
            req = urllib.request.Request(api_url, data=payload,
                headers={"Authorization":f"Bearer {api_key}","Content-Type":"application/json"})
            resp = json.loads(urllib.request.urlopen(req, timeout=90, context=ctx).read())
            text = resp["choices"][0]["message"]["content"].strip()
            parts = text.split("===NEXT===")
            for j, (aid, _) in enumerate(chunk):
                results[aid] = parts[j].strip() if j < len(parts) else ""
        except:
            for aid, _ in chunk: results[aid] = ""
        time.sleep(0.5)
    return results

def build(data):
    arts = [a for a in data["articles"] if a["source"] not in ("IRNA (Iran)", "Tanjug (Serbia)")]
    
    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(21); s.page_height = Cm(29.7)
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)
    
    style = doc.styles['Normal']
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(4)
    
    # Cover
    for _ in range(4): doc.add_paragraph("")
    bar = doc.add_paragraph(); bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = bar.add_run(" "*100); r.font.size = Pt(2)
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"),"C41E3A"); shd.set(qn("w:val"),"clear")
    r._element.get_or_add_rPr().append(shd)
    doc.add_paragraph("")
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("参 考 消 息"); r.font.size = Pt(34); r.bold = True
    r.font.color.rgb = RGBColor(0xC4,0x1E,0x3A)
    doc.add_paragraph("")
    t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t2.add_run("国外视角下的中国"); r2.font.size = Pt(18)
    r2.font.color.rgb = RGBColor(0x44,0x44,0x44)
    doc.add_paragraph("")
    b2 = doc.add_paragraph(); b2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = b2.add_run(" "*100); r.font.size = Pt(1)
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"),"C41E3A"); shd.set(qn("w:val"),"clear")
    r._element.get_or_add_rPr().append(shd)
    for _ in range(3): doc.add_paragraph("")
    
    now = datetime.now()
    weekdays = ["Monday","Tuesday","Wednesday","Thursday","Friday","Saturday","Sunday"]
    wd = weekdays[now.weekday()]
    dp = doc.add_paragraph(); dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dp.add_run(f"  {now.year}-{now.month:02d}-{now.day:02d}  {wd}  "); dr.font.size = Pt(13)
    dr.font.color.rgb = RGBColor(0x66,0x66,0x66)
    doc.add_paragraph("")
    
    ft_cnt = sum(1 for a in arts if a.get("full_text") and len(a["full_text"])>80)
    sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sp.add_run(f"{len(arts)} articles · {len(SRC_META)} sources · {ft_cnt} full text")
    r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x99,0x99,0x99)
    
    src_set = sorted(set(SRC_FLAGS[s] for s in set(a["source"] for a in arts) if s in SRC_FLAGS))
    fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("  ".join(src_set))
    fr.font.size = Pt(7); fr.font.color.rgb = RGBColor(0xAA,0xAA,0xAA)
    
    doc.add_page_break()
    
    # Translate
    print("Translating titles...", flush=True)
    all_titles = [a["title"] for a in arts]
    cn_titles = translate(all_titles)
    print(f"  Titles: {sum(1 for t in cn_titles if t)}/{len(cn_titles)}", flush=True)
    
    print("Translating full texts...", flush=True)
    ft_cn_map = translate_full(arts)
    print(f"  Full texts: {sum(1 for v in ft_cn_map.values() if v)}/{len(ft_cn_map)}", flush=True)
    
    groups = {}
    for a in arts: groups.setdefault(a["source"],[]).append(a)
    
    tt_idx = 0; gidx = 0
    for src, items in groups.items():
        color, full_name = SRC_META.get(src, ("333333", src))
        color_rgb = RGBColor.from_string(color)
        add_section_header(doc, f"{full_name}  |  {len(items)} articles", color)
        
        for item in items:
            gidx += 1
            cn_title = cn_titles[tt_idx] if tt_idx < len(cn_titles) else ""; tt_idx += 1
            cn_full = ft_cn_map.get(id(item), "")
            
            tp = doc.add_paragraph()
            tp.paragraph_format.space_before = Pt(10)
            tp.paragraph_format.space_after = Pt(1)
            tp.paragraph_format.line_spacing = 1.3
            
            nr = tp.add_run(f" {gidx:03d} "); nr.font.size = Pt(7); nr.bold = True
            nr.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), color); shd.set(qn("w:val"),"clear")
            nr._element.get_or_add_rPr().append(shd)
            
            title = item['title'].replace("&lt;","<").replace("&gt;",">").replace("&amp;","&")
            tr = tp.add_run(f"  {title}"); tr.font.size = Pt(11); tr.bold = True
            tr.font.color.rgb = RGBColor(0x1A,0x1A,0x1A)
            
            if cn_title:
                ct = doc.add_paragraph()
                ct.paragraph_format.space_before = Pt(0); ct.paragraph_format.space_after = Pt(3)
                ct.paragraph_format.left_indent = Cm(0.8); ct.paragraph_format.line_spacing = 1.2
                cr = ct.add_run(cn_title); cr.font.size = Pt(9.5)
                cr.font.color.rgb = RGBColor(0x66,0x66,0x66); cr.italic = True
            
            sm = item.get("summary","")
            if sm:
                sm = sm.replace("&lt;","<").replace("&gt;",">").replace("&amp;","&").replace("&#x27;","'").replace("&quot;",'"')
            if sm and len(sm) > 15:
                sp = doc.add_paragraph()
                sp.paragraph_format.space_before = Pt(3); sp.paragraph_format.space_after = Pt(2)
                sp.paragraph_format.left_indent = Cm(0.8); sp.paragraph_format.line_spacing = 1.2
                lb = sp.add_run("Why it matters  "); lb.font.size = Pt(7); lb.bold = True
                lb.font.color.rgb = color_rgb
                sr = sp.add_run(sm[:400]); sr.font.size = Pt(9)
                sr.font.color.rgb = RGBColor(0x55,0x55,0x55)
            
            ft_text = item.get("full_text","")
            if ft_text:
                ft_text = ft_text.replace("&lt;","<").replace("&gt;",">").replace("&amp;","&").replace("&#x27;","'").replace("&quot;",'"')
            if ft_text and len(ft_text) > 80:
                fp = doc.add_paragraph()
                fp.paragraph_format.space_before = Pt(4); fp.paragraph_format.left_indent = Cm(0.8)
                fp.paragraph_format.line_spacing = 1.3
                lb = fp.add_run("EN  "); lb.font.size = Pt(7); lb.bold = True
                lb.font.color.rgb = RGBColor(0xAA,0xAA,0xAA)
                fr = fp.add_run(ft_text[:1000]); fr.font.size = Pt(8.5)
                fr.font.color.rgb = RGBColor(0x44,0x44,0x44)
                if cn_full:
                    fcp = doc.add_paragraph()
                    fcp.paragraph_format.space_before = Pt(1); fcp.paragraph_format.left_indent = Cm(0.8)
                    fcp.paragraph_format.line_spacing = 1.3
                    lb2 = fcp.add_run("CN  "); lb2.font.size = Pt(7); lb2.bold = True
                    lb2.font.color.rgb = color_rgb
                    fr2 = fcp.add_run(cn_full); fr2.font.size = Pt(8.5)
                    fr2.font.color.rgb = RGBColor(0x33,0x33,0x33)
            
            if item != items[-1]:
                sp2 = doc.add_paragraph()
                sp2.paragraph_format.space_before = Pt(6); sp2.paragraph_format.space_after = Pt(0)
                dr2 = sp2.add_run("─" * 40); dr2.font.size = Pt(6)
                dr2.font.color.rgb = RGBColor(0xDD,0xDD,0xDD)
    
    doc.add_paragraph("")
    add_divider(doc, "C41E3A", 6)
    fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("参考消息 · 国外视角下的中国 · 每日自动生成")
    fr.font.size = Pt(7); fr.font.color.rgb = RGBColor(0xAA,0xAA,0xAA)
    
    return doc

if __name__ == "__main__":
    inp = sys.argv[1] if len(sys.argv)>1 else "china_news_raw.json"
    with open(inp,"r",encoding="utf-8") as f:
        data = json.load(f)
    print(f"Building doc ({data['total']} articles)...", flush=True)
    doc = build(data)
    now = datetime.now()
    out = f"{now.year}年{now.month}月{now.day}日_参考消息-国外视角下的中国.docx"
    doc.save(out)
    print(f"Saved: {out}", flush=True)
